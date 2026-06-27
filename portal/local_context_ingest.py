"""FR-013a P2 — ingest a document into candidate local-context items.

Pipeline: extract text per format (PDF / Word / Excel / plain) → structure it into
discrete TYPED items (standing_order / medication / treatment_priority) → add them
**INACTIVE** (source = filename) for review/edit/activate in the P3 UI. Nothing goes
live until the instructor confirms (FR-008 "nothing staged until you confirm").

Structuring uses Claude when a key is available (high fidelity — handles tables /
formularies / prose); a heuristic line-split is the no-key / failure fallback. The
spec flags parsing fidelity as the risk, so we lean on the P3 review step: rough
candidates are fine, the instructor fixes type/title/content before activating.

extract_text / parse_items are pure + offline-testable; only the Claude path in
parse_items touches the network (and degrades to the heuristic on any failure).
"""
from __future__ import annotations

import io
import re
from typing import Any

from . import local_context as _lc
from . import scenario_gen as _sg  # reuse GEN_MODEL + extract_json

MAX_TEXT = 20000   # chars of document text sent to the model (bound cost/latency)
MAX_ITEMS = 60     # cap candidates added per upload
SUPPORTED = ("pdf", "docx", "xlsx", "xlsm", "txt", "md", "csv", "tsv")


def _ext(filename: str) -> str:
    return (filename.rsplit(".", 1)[-1] if "." in (filename or "") else "").lower()


def extract_text(filename: str, data: bytes) -> str:
    """Pull plain text out of an uploaded document. Per-format libs are imported
    lazily so a portal without them still runs (the format just errors clearly)."""
    ext = _ext(filename)
    if ext == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:  # pragma: no cover
            raise ValueError("PDF support not installed (pip install pypdf)")
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if ext == "docx":
        try:
            import docx
        except ImportError:  # pragma: no cover
            raise ValueError("Word support not installed (pip install python-docx)")
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext in ("xlsx", "xlsm"):
        try:
            import openpyxl
        except ImportError:  # pragma: no cover
            raise ValueError("Excel support not installed (pip install openpyxl)")
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c not in (None, "")]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if ext in ("txt", "md", "csv", "tsv"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"unsupported file type: .{ext or '?'} (supported: {', '.join(SUPPORTED)})")


def _heuristic_items(text: str) -> list[dict[str, Any]]:
    """No-AI fallback: each substantive line/bullet → a candidate standing order.
    The instructor re-types the type/title in review; this just gets text in."""
    items: list[dict[str, Any]] = []
    for ln in text.splitlines():
        s = re.sub(r"^[\-\*•\d.)\s]+", "", ln).strip()
        if len(s) < 4:
            continue
        title = (s[:60] + "…") if len(s) > 60 else s
        items.append({"type": "standing_order", "title": title, "content": s})
        if len(items) >= MAX_ITEMS:
            break
    return items


def _ai_items(text: str, api_key: str) -> list[dict[str, Any]]:
    from anthropic import Anthropic
    system = (
        "You extract A SITE'S LOCAL clinical rules from a document into discrete "
        "items for a simulation's local-practice overlay. Each item is exactly one "
        "of: standing_order, medication (local formulary entry), treatment_priority. "
        "Skip headers, page numbers, and boilerplate; keep concrete, actionable rules. "
        'Return STRICT JSON only (no prose): {"items":[{"type":"standing_order|'
        'medication|treatment_priority","title":"short label","content":"the rule"}]}'
    )
    client = Anthropic(api_key=api_key)
    resp = client.messages.create(
        model=_sg.GEN_MODEL, max_tokens=_sg.MAX_TOKENS, system=system,
        messages=[{"role": "user", "content": text[:MAX_TEXT]}],
    )
    out = "".join(getattr(b, "text", "") for b in (resp.content or []))
    obj = _sg.extract_json(out)            # tolerant object extraction
    raw = obj.get("items") if isinstance(obj, dict) else None
    return raw if isinstance(raw, list) else []


def parse_items(text: str, *, api_key: str = "") -> list[dict[str, Any]]:
    """Document text → normalized candidate items. Claude first (if key), heuristic
    fallback; every item is coerced to a valid type with a non-empty title."""
    text = (text or "").strip()
    if not text:
        return []
    items: list[dict[str, Any]] = []
    if api_key:
        try:
            items = _ai_items(text, api_key)
        except Exception:  # noqa: BLE001 — any AI/parse failure → heuristic
            items = []
    if not items:
        items = _heuristic_items(text)
    out: list[dict[str, Any]] = []
    for it in items:
        if not isinstance(it, dict):
            continue
        content = str(it.get("content") or "").strip()
        if not content:
            continue
        t = str(it.get("type") or "").strip()
        if t not in _lc.ITEM_TYPES:
            t = "standing_order"
        title = str(it.get("title") or "").strip() or (
            (content[:60] + "…") if len(content) > 60 else content)
        out.append({"type": t, "title": title, "content": content})
        if len(out) >= MAX_ITEMS:
            break
    return out


def ingest(filename: str, data: bytes, *, api_key: str = "") -> dict[str, Any]:
    """Extract → parse → add candidates INACTIVE (source=filename). Returns
    {added, items}. Raises ValueError for an unsupported / unreadable file."""
    text = extract_text(filename, data)
    added = [
        _lc.add_item(type=c["type"], title=c["title"], content=c["content"],
                     source=filename or "upload", active=False)
        for c in parse_items(text, api_key=api_key)
    ]
    return {"added": len(added), "items": added}
